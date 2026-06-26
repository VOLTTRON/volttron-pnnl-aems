"""Build the AEMS Software Deployment Guide .docx from Markdown source.

This builder clones an existing PNNL-styled .docx (the Installer Guide) to
inherit its full house style — cover page layout, headers/footers, section
properties, PNNL_Title_Page_Title / PNNL_Title-Page_Text / Heading-Front (TOC)
/ Body Text / HTML Code / Caption-Fig / Caption-Tab / Acronyms styles, page
numbering, and margins — then strips the template's body content and rebuilds
the body from the Markdown source with PNNL named styles applied.

Run from the repository root:
    python docs/proposed/aems-deployment-guide/pandoc/md_to_docx.py
"""

from __future__ import annotations

import copy
import re
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[4]
SRC = REPO_ROOT / "docs" / "proposed" / "aems-deployment-guide" / "aems-deployment-guide.md"
TEMPLATE = REPO_ROOT / "AEMS Building Installer Configuration User Guide (2025-06-26).docx"
OUT_NAME = f"AEMS Software Deployment Guide ({date.today().isoformat()}).docx"
OUT = REPO_ROOT / OUT_NAME

# Map our content "kinds" to PNNL named styles in the template.
PNNL_STYLES = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "front_h1": "Heading-Front (TOC)",
    "front_h1_notoc": "Heading Front (No TOC)",
    "body": "Body Text",
    "code": "HTML Code",
    "code_block": "HTML Preformatted",
    "caption_fig": "Caption-Fig",
    "caption_tab": "Caption-Tab",
    "acronym": "Acronyms",
    "table_text": "TableText",
    "bullet": "List Bullet",
    "number": "List Number",
}

# Headings that should use the front-matter style (no number, used by ToC)
FRONT_HEADINGS = {
    "preface",
    "audience and scope",
    "how to read this guide",
    "document conventions",
    "acronyms",
}


# -------- inline span parser --------------------------------------------

# Order matters: more-specific patterns first. Autolinks (<https://...>) must
# match before angle-bracket placeholders like <HOSTNAME>. The § reference
# pattern (§N, §N.N, §N.N.N) is resolved against a section-number → bookmark
# index populated during the heading walk in build().
INLINE_RE = re.compile(
    r"(`[^`]+`)"
    r"|(\*\*[^*]+\*\*)"
    r"|(\*[^*]+\*)"
    r"|(\[[^\]]+\]\([^)]+\))"
    r"|(<https?://[^>\s]+>)"
    r"|(§\d+(?:\.\d+)*)"
)

# Index of section-number string ("4", "4.2", "13.9") to bookmark name.
# Populated during the pre-walk in build() before any inline rendering.
SECTION_BOOKMARKS = {}

# Base URL for in-repo file references. The Markdown source uses relative
# paths like '../../../aems-app/README.md' so that an editor previewing
# the Markdown can resolve the link locally; when emitting the .docx we
# rewrite those into GitHub blob/tree URLs against this base so the
# hyperlinks are clickable in the published document.
GITHUB_REPO = "https://github.com/VOLTTRON/volttron-pnnl-aems"
GITHUB_BRANCH = "main"

# The Markdown file lives at
#   docs/proposed/aems-deployment-guide/aems-deployment-guide.md
# so '../../../FOO' (three '..' segments) resolves to FOO at the repo
# root, '../../../aems-app/README.md' resolves to aems-app/README.md, etc.
# Anything that doesn't already start with http(s) and isn't an anchor
# is treated as a repo-relative path.
_MARKDOWN_DIR_FROM_REPO = "docs/proposed/aems-deployment-guide"


def _resolve_repo_path(url):
    """Resolve a Markdown-relative link target to an absolute path within
    the repository (no leading slash, no scheme). Returns None for purely
    in-page anchor links or mailto links.
    """
    if not url:
        return None
    if url.startswith("#") or url.startswith("mailto:") or url.startswith("tel:"):
        return None
    if url.startswith("http://") or url.startswith("https://"):
        return None
    # Strip any trailing query/fragment for the path resolution
    path_part, _, suffix = url.partition("#")
    path_part, _, query = path_part.partition("?")
    # Build the absolute repo-relative path by joining and normalizing
    parts = (_MARKDOWN_DIR_FROM_REPO + "/" + path_part).split("/")
    stack = []
    for part in parts:
        if part == "" or part == ".":
            continue
        if part == "..":
            if stack:
                stack.pop()
            continue
        stack.append(part)
    resolved = "/".join(stack)
    return (resolved, suffix, query)


def _to_github_url(url):
    """Convert a Markdown-relative repo path into a GitHub URL. Directories
    use the 'tree' base; files use 'blob'. URLs that already point at
    http(s) targets are returned unchanged.
    """
    if url.startswith("http://") or url.startswith("https://"):
        return url
    parts = _resolve_repo_path(url)
    if parts is None:
        return url
    resolved, suffix, query = parts
    if not resolved:
        # Resolved to the repo root itself
        target = f"{GITHUB_REPO}/tree/{GITHUB_BRANCH}"
    else:
        # Directory links in the Markdown source end with '/'; files don't.
        is_dir = url.rstrip("?#").endswith("/")
        kind = "tree" if is_dir else "blob"
        target = f"{GITHUB_REPO}/{kind}/{GITHUB_BRANCH}/{resolved}"
    if query:
        target += "?" + query
    if suffix:
        target += "#" + suffix
    return target


# Reusable relationship type for external hyperlinks
_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


def _add_hyperlink(paragraph, text, url):
    """Insert a clickable external hyperlink into *paragraph*. Creates a
    relationship in the document part so Word renders the run as a real
    link, not just blue-underlined text.
    """
    part = paragraph.part
    r_id = part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    _append_hyperlink_run(hyperlink, text)
    paragraph._p.append(hyperlink)


def _add_internal_link(paragraph, text, bookmark_name):
    """Insert a clickable cross-reference to a bookmark in this document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    _append_hyperlink_run(hyperlink, text)
    paragraph._p.append(hyperlink)


def _append_hyperlink_run(hyperlink_el, text):
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B4F9C")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink_el.append(new_run)


def _bookmark_for_section(num_str):
    """Bookmark name for a section number like '4', '4.2', '13.9'. Word
    bookmark names cannot contain '.', so we replace dot with underscore.
    """
    return "sec_" + num_str.replace(".", "_")


def add_inline(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("[") and "](" in token:
            label, _, rest = token[1:].partition("](")
            url = rest[:-1]
            # External http(s) URLs and same-repo relative paths both
            # benefit from being real Word hyperlinks. Skip purely
            # anchor-only links (#section) to avoid creating broken
            # external relationships. Relative paths to in-repo files
            # are rewritten as GitHub blob/tree URLs so the link works
            # when the .docx is read off-host.
            if url.startswith(("http://", "https://")):
                _add_hyperlink(paragraph, label, url)
            elif url.startswith("#"):
                # In-document anchor — keep as plain styled text rather
                # than emit a broken external relationship. Anchor-style
                # navigation is handled by the §N cross-references.
                run = paragraph.add_run(label)
                run.font.color.rgb = RGBColor(0x0B, 0x4F, 0x9C)
                run.underline = True
            elif "/" in url or "." in url:
                gh_url = _to_github_url(url)
                _add_hyperlink(paragraph, label, gh_url)
            else:
                run = paragraph.add_run(label)
                run.font.color.rgb = RGBColor(0x0B, 0x4F, 0x9C)
                run.underline = True
        elif token.startswith("<http"):
            url = token[1:-1]
            _add_hyperlink(paragraph, url, url)
        elif token.startswith("§"):
            num = token[1:]
            bookmark = SECTION_BOOKMARKS.get(num)
            if bookmark:
                _add_internal_link(paragraph, token, bookmark)
            else:
                # Unknown section reference — keep the literal text so the
                # reviewer can see and fix it. Print a warning for debug.
                print(f"warning: unresolved section reference {token!r}", file=sys.stderr)
                paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# -------- frontmatter ---------------------------------------------------


def parse_frontmatter(lines):
    if not lines or lines[0].strip() != "---":
        return {}, lines
    meta = {}
    i = 1
    while i < len(lines) and lines[i].strip() != "---":
        line = lines[i].rstrip()
        if line.startswith("  "):
            pass
        elif ":" in line:
            key, _, value = line.partition(":")
            meta[key.strip()] = value.strip().strip('"')
        i += 1
    return meta, lines[i + 1 :]


# -------- block tokenizer -----------------------------------------------


def parse_blocks(lines):
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        m = re.match(r"^<!--\s*diagram:\s*(\S+)\s*-->$", stripped)
        if m:
            yield ("diagram", m.group(1))
            i += 1
            continue

        if stripped == "---":
            yield ("hr", None)
            i += 1
            continue

        if stripped.startswith("```"):
            j = i + 1
            code = []
            while j < n and not lines[j].strip().startswith("```"):
                code.append(lines[j])
                j += 1
            yield ("code", "\n".join(code))
            i = j + 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            yield (f"h{min(level, 4)}", text)
            i += 1
            continue

        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows = []
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(row)
                j += 1
            yield ("table", (header, rows))
            i = j
            continue

        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).rstrip())
                i += 1
                while i < n and lines[i].startswith(("    ", "\t")) and lines[i].strip():
                    items[-1] += " " + lines[i].strip()
                    i += 1
            yield ("ul", items)
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < n and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]).rstrip())
                i += 1
                while i < n and lines[i].startswith(("    ", "\t")) and lines[i].strip():
                    items[-1] += " " + lines[i].strip()
                    i += 1
            yield ("ol", items)
            continue

        if stripped.startswith(">"):
            quote = []
            while i < n and lines[i].lstrip().startswith(">"):
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]).rstrip())
                i += 1
            yield ("blockquote", "\n".join(quote))
            continue

        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _starts_new_block(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        yield ("p", " ".join(para_lines))


def _starts_new_block(line):
    s = line.strip()
    return (
        s.startswith("#")
        or s.startswith("```")
        or s.startswith(">")
        or s == "---"
        or re.match(r"^\s*[-*]\s+", line) is not None
        or re.match(r"^\s*\d+\.\s+", line) is not None
        or ("|" in s and s.count("|") >= 2)
    )


# -------- template handling ---------------------------------------------


def clone_template_and_clear_body(meta):
    """Copy the template .docx to OUT, then strip its body paragraphs after the
    title-page section so we can append our content while keeping styles,
    headers/footers, section properties, and cover graphics.
    """
    if not TEMPLATE.exists():
        print(
            f"ERROR: Template not found at {TEMPLATE}.\n"
            "       The builder requires the Installer Guide .docx as a style source.",
            file=sys.stderr,
        )
        sys.exit(1)
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(str(OUT))

    body = doc.element.body

    # Identify the sectPr at the very end (last sectPr is the body's section
    # properties, e.g. margins for the entire document). Preserve it.
    sect_prs = body.findall(qn("w:sectPr"))
    last_sectPr = body.find(qn("w:sectPr"))

    # Find the index of the FIRST content paragraph after the title page.
    # The title page in PNNL guides ends with a paragraph that has a
    # sectPr embedded in its pPr (section break). Locate the first
    # paragraph whose style is *not* PNNL title-page related and whose
    # text is non-empty — that's where original content starts.
    children = list(body)
    first_content_idx = None
    for i, el in enumerate(children):
        tag = el.tag.split("}")[-1]
        if tag != "p":
            continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t")))
        if not text.strip():
            continue
        # Get pStyle val
        pPr = el.find(qn("w:pPr"))
        style_val = ""
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_val = pStyle.get(qn("w:val")) or ""
        if style_val in (
            "PNNL_Title_Page_Title",
            "PNNL_Title-Page_Text",
            "White_Cover_PNNL_Prepared_For",
        ):
            continue
        # First real content paragraph
        first_content_idx = i
        break

    if first_content_idx is None:
        print("ERROR: Could not locate content start in template.", file=sys.stderr)
        sys.exit(1)

    # Capture the section-break paragraph that separates title page from
    # body (the paragraph just before first_content_idx that contains a
    # sectPr in its pPr). We need to keep title-page section break.
    # Then remove everything from first_content_idx onward, but preserve
    # the final body sectPr.
    saved_final_sectPr = None
    if last_sectPr is not None:
        saved_final_sectPr = copy.deepcopy(last_sectPr)

    # Remove all body children from first_content_idx onward
    for el in children[first_content_idx:]:
        body.remove(el)

    # Re-attach the final sectPr so page geometry of the body persists
    if saved_final_sectPr is not None:
        body.append(saved_final_sectPr)

    # Override title-page text: find any paragraph with PNNL_Title-Page_Text
    # containing the old title (e.g. "AEMS Building Installer Configuration
    # User Guide") and replace its text with our title. Also handle "April
    # 2026" / report number / authors.
    _patch_title_page(doc, meta)

    return doc


_MONTHS = (
    "January|February|March|April|May|June|July|"
    "August|September|October|November|December"
)


def _patch_title_page(doc, meta):
    """Replace title-page text in the template. Cover content lives inside
    nested <w:sdt> (Structured Document Tag) blocks in the PNNL template,
    so we recurse to every <w:p> and match on its pStyle value.
    """
    title = meta.get("title", "AEMS Software Deployment Guide")
    subtitle = meta.get("subtitle", "")
    pub_date = meta.get("date", "June 2026")
    report_number = meta.get("report_number", "PNNL-XXXXX")
    contract = meta.get(
        "contract", "Prepared for the U.S. Department of Energy under Contract DE-AC05-76RL01830"
    )
    # Authors may be a literal "\n" separator in the YAML; render as
    # newline-separated runs on the cover.
    authors_raw = meta.get("authors", "Pacific Northwest National Laboratory")
    authors = authors_raw.replace("\\n", "\n")

    # PNNL named styles used on the cover. Map each to the replacement text.
    style_replacements = {
        "PNNLCoverNumber": report_number,
        "PNNLCoverTitle": title,
        "PNNLCoverDate": pub_date,
        "PNNLCoverAuthors": authors,
        "PNNLCoverContract": contract,
        "PNNLTitlePageTitle": title,
    }

    body = doc.element.body
    title_paras = []
    for p_el in body.iter(qn("w:p")):
        pPr = p_el.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        style_val = pStyle.get(qn("w:val")) or ""
        if style_val in style_replacements:
            _replace_paragraph_text(p_el, style_replacements[style_val])
            if style_val == "PNNLCoverTitle":
                title_paras.append(p_el)

    # Insert subtitle paragraph just after each cover-title paragraph
    if subtitle:
        for tp in title_paras:
            new_p = OxmlElement("w:p")
            npPr = OxmlElement("w:pPr")
            npStyle = OxmlElement("w:pStyle")
            # Reuse the PNNLCoverDate style for the subtitle so spacing /
            # alignment / color match the rest of the cover block; the
            # subtitle is rendered italic via run formatting below.
            npStyle.set(qn("w:val"), "PNNLCoverDate")
            npPr.append(npStyle)
            new_p.append(npPr)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            it = OxmlElement("w:i")
            rPr.append(it)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = subtitle
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            new_p.append(r)
            tp.addnext(new_p)

    # Also handle ordinary (non-SDT) paragraphs that still carry the old
    # title's text — belt and braces in case the template version differs.
    for p_el in body.iter(qn("w:p")):
        text = "".join(t.text or "" for t in p_el.iter(qn("w:t")))
        if not text.strip():
            continue
        if "Installer Configuration" in text or "Owner and Occupant" in text:
            _replace_paragraph_text(p_el, title)
        elif re.match(rf"^({_MONTHS})\s+\d{{4}}\s*$", text.strip()) and "March 2024" in text:
            _replace_paragraph_text(p_el, pub_date)


def _replace_paragraph_text(p_el, new_text):
    """Replace the text content of a paragraph element while keeping its
    first run's formatting. Removes other runs.

    PNNL cover paragraphs nest their visible content inside a structured
    document tag (<w:sdt>/<w:sdtContent>/<w:r>), so direct .findall(w:r) on
    the paragraph returns nothing. We collect all <w:r> descendants of the
    paragraph (including those inside SDT content blocks) and keep the
    first one's formatting.

    If *new_text* contains newlines, each line becomes a separate text
    chunk inside the first run, separated by w:br elements — Word
    renders that as a line break within the paragraph, which is what the
    PNNL_CoverAuthors style expects (one author per line, same style).
    """
    lines = new_text.split("\n")
    runs = list(p_el.iter(qn("w:r")))
    if not runs:
        # No runs anywhere — synthesize one and append to paragraph
        new_r = OxmlElement("w:r")
        for idx, line in enumerate(lines):
            if idx:
                new_r.append(OxmlElement("w:br"))
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = line
            new_r.append(t)
        p_el.append(new_r)
        return

    first_r = runs[0]
    # Replace first run's text + br + tab children with the new lines
    for tag in ("w:t", "w:tab", "w:br"):
        for el in first_r.findall(qn(tag)):
            first_r.remove(el)
    for idx, line in enumerate(lines):
        if idx:
            first_r.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        first_r.append(t)

    # Remove all other runs (in-paragraph or SDT-wrapped). Each run's
    # parent might be sdtContent rather than the paragraph itself.
    for r in runs[1:]:
        parent = r.getparent()
        if parent is not None:
            parent.remove(r)


# -------- rendering -----------------------------------------------------


def add_paragraph_with_style(doc, text, style_name):
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
    if text:
        add_inline(p, text)
    return p


# Running section counters maintained by render_heading. Reset at the
# start of each build() call.
_SECTION_STATE = {
    "h1": 0,
    "h2": 0,
    "h3": 0,
    "bookmark_id": 1000,
    "in_unnumbered_h1": False,
}


def _next_bookmark_id():
    _SECTION_STATE["bookmark_id"] += 1
    return _SECTION_STATE["bookmark_id"]


def _add_bookmark_around_paragraph(p, name):
    """Wrap a paragraph's runs in matching bookmarkStart / bookmarkEnd tags
    so the entire heading is the link target. bookmarkStart must come
    before the first run; bookmarkEnd after the last run.
    """
    bm_id = _next_bookmark_id()
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))
    # Insert bookmarkStart as the first child after pPr; bookmarkEnd at end
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bm_start)
    else:
        p._p.insert(0, bm_start)
    p._p.append(bm_end)


def render_heading(doc, level, text):
    """Render a heading and, for numbered sections, register a bookmark so
    inline §N / §N.N references can link to it.
    """
    norm = text.strip().lower()

    # Front-matter and appendix headings don't get a number, but should
    # still get bookmarks named after a slugified form of their title in
    # case we ever want to link to them.
    if level == 1 and norm in FRONT_HEADINGS:
        _SECTION_STATE["in_unnumbered_h1"] = True
        p = doc.add_paragraph(style=PNNL_STYLES["front_h1"])
        p.add_run(text)
        _add_bookmark_around_paragraph(p, "fm_" + _slug(text))
        return
    if level == 1 and norm.startswith("appendix "):
        _SECTION_STATE["in_unnumbered_h1"] = True
        p = doc.add_paragraph(style=PNNL_STYLES["front_h1_notoc"])
        p.add_run(text)
        _add_bookmark_around_paragraph(p, "app_" + _slug(text))
        return
    # Sub-headings inside an unnumbered H1 (front matter, appendix) inherit
    # no number; render them with the same outline level but skip the
    # numbering+bookmark assignment.
    if _SECTION_STATE["in_unnumbered_h1"] and level >= 2:
        p = doc.add_paragraph(style=f"Heading {level}")
        p.add_run(text)
        return

    # Numbered section path
    if level == 1:
        _SECTION_STATE["in_unnumbered_h1"] = False
        _SECTION_STATE["h1"] += 1
        _SECTION_STATE["h2"] = 0
        _SECTION_STATE["h3"] = 0
        num_str = str(_SECTION_STATE["h1"])
    elif level == 2:
        _SECTION_STATE["h2"] += 1
        _SECTION_STATE["h3"] = 0
        num_str = f'{_SECTION_STATE["h1"]}.{_SECTION_STATE["h2"]}'
    elif level == 3:
        _SECTION_STATE["h3"] += 1
        num_str = f'{_SECTION_STATE["h1"]}.{_SECTION_STATE["h2"]}.{_SECTION_STATE["h3"]}'
    else:
        num_str = None

    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    if num_str:
        bookmark_name = _bookmark_for_section(num_str)
        SECTION_BOOKMARKS[num_str] = bookmark_name
        _add_bookmark_around_paragraph(p, bookmark_name)


def _slug(text):
    s = re.sub(r"[^A-Za-z0-9]+", "_", text.strip().lower())
    return s.strip("_") or "section"


def render_body(doc, text):
    add_paragraph_with_style(doc, text, PNNL_STYLES["body"])


def _force_full_width(tbl):
    """Apply 100% page-width and autofit layout to a python-docx table.
    Mirrors the logic in render_table so callouts and code blocks match
    the data tables' full-width behaviour.
    """
    tblPr = tbl._tbl.tblPr
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)
    # Suppress the firstRow/lastRow style hooks that would otherwise pull
    # in the PNNL header banding from any table style applied.
    existing_look = tblPr.find(qn("w:tblLook"))
    if existing_look is not None:
        tblPr.remove(existing_look)
    look = OxmlElement("w:tblLook")
    look.set(qn("w:val"), "0000")
    look.set(qn("w:firstRow"), "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "1")
    look.set(qn("w:noVBand"), "1")
    tblPr.append(look)


def _set_cell_borders(cell, color_hex, sz="6"):
    """Set all four borders of a cell to a single color / weight."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def render_code(doc, code_text):
    """Render a fenced code block as a full-width, single-cell table with a
    black background and white monospace text — a terminal/console look.
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    _force_full_width(tbl)

    cell = tbl.rows[0].cells[0]
    # Black background, dark border so the box reads as a console
    _shade_cell(cell, "111111")
    _set_cell_borders(cell, "111111", sz="4")

    # Clear the default paragraph and emit one per code line so wrapping
    # behaves like a real terminal — long lines wrap, blank lines remain
    # blank, indentation preserved by no-break spaces.
    cell.text = ""
    lines = code_text.splitlines() or [""]
    for idx, line in enumerate(lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        # Apply Normal style so we control formatting via run properties
        # rather than inheriting body-text spacing from Body Text.
        try:
            p.style = doc.styles["Normal"]
        except KeyError:
            pass
        # Tight line spacing for code
        pPr = p._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is not None:
            pPr.remove(spacing)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "260")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

        # Preserve leading whitespace explicitly. add_run does set
        # xml:space=preserve on the text element, but Word still collapses
        # multiple spaces in display unless we use it; convert tabs to
        # 4 spaces for predictable rendering.
        rendered = line.replace("\t", "    ")
        run = p.add_run(rendered if rendered else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xF2, 0xF2, 0xF2)
    doc.add_paragraph()


def render_table(doc, header, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Apply the PNNL house table style. PNNL_Solid_Header_Row gives an
    # orange (D57500) header row with white bold text and gray banded
    # rows — matches the existing Installer and Owner guides.
    pnnl_table_style = None
    for candidate in ("PNNL_Solid_Header_Row", "PNNL_Single_Header_Row_Footnote"):
        try:
            tbl.style = doc.styles[candidate]
            pnnl_table_style = candidate
            break
        except KeyError:
            continue
    if pnnl_table_style is None:
        try:
            tbl.style = "Table Grid"
        except KeyError:
            pass

    # Enable the firstRow / horizontal banding conditional formatting on
    # this table instance. python-docx does not expose tblLook, so we set
    # it directly via OXML. Without this, Word does not apply the style's
    # firstRow run/cell formatting and the header row stays unstyled.
    tblPr = tbl._tbl.tblPr
    existing_look = tblPr.find(qn("w:tblLook"))
    if existing_look is not None:
        tblPr.remove(existing_look)
    look = OxmlElement("w:tblLook")
    look.set(qn("w:val"), "04A0")
    look.set(qn("w:firstRow"), "1")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "0")
    look.set(qn("w:noVBand"), "1")
    tblPr.append(look)

    # Force full page width. Without an explicit tblW Word lays out the
    # table at content-natural width, so short headers like "Convention"
    # leave acres of empty page beside the table. 5000/pct = 100% of the
    # available text area between the page margins.
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    # Distribute columns evenly. tblLayout=autofit lets Word resize
    # individual columns to fit content while still respecting the full
    # 100% table width above.
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)

    hdr_cells = tbl.rows[0].cells
    for j, h in enumerate(header):
        hdr_cells[j].text = ""
        p = hdr_cells[j].paragraphs[0]
        try:
            p.style = doc.styles[PNNL_STYLES["table_text"]]
        except KeyError:
            pass
        # The PNNL table style supplies the orange fill, white color, and
        # bold weight via firstRow conditional formatting — do NOT force
        # those properties here, or the manual run color (#FFFFFF) ends
        # up overriding the style and looks identical regardless of the
        # actual table-style fill the reader sees.
        run = p.add_run(h)
        if pnnl_table_style is None:
            # Fallback: no PNNL style available; render a neutral bold header.
            run.bold = True

    for i, row in enumerate(rows, start=1):
        for j, cell_text in enumerate(row):
            if j >= len(header):
                break
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            try:
                p.style = doc.styles[PNNL_STYLES["table_text"]]
            except KeyError:
                pass
            add_inline(p, cell_text)
    doc.add_paragraph()


def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def render_list(doc, items, kind):
    style = PNNL_STYLES["bullet"] if kind == "ul" else PNNL_STYLES["number"]
    for item in items:
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
            p.add_run("• " if kind == "ul" else "")
        add_inline(p, item)


# Callout color palette. Each entry: (header_fill, body_fill, body_border)
_CALLOUT_PALETTE = {
    "WARNING": ("C0392B", "FDEDEC", "C0392B"),  # red header / pink body
    "NOTE":    ("0B4F9C", "EAF2FB", "0B4F9C"),  # blue header / pale-blue body
    "TIP":     ("1E8449", "EAF6EE", "1E8449"),  # green header / pale-green body
}


def render_blockquote(doc, text):
    """Render NOTE / TIP / WARNING callouts as a two-row, one-column,
    full-width table. The header row is a colored band carrying the
    label in bold white; the body row holds the prose with a faint tint
    of the same hue and matching border. Plain blockquotes fall back to
    a single-cell gray box.
    """
    t = text.strip()
    label = None
    for kind in _CALLOUT_PALETTE:
        if t.startswith(f"**{kind}."):
            label = kind
            break

    if label is None:
        # Plain blockquote — single-cell gray box.
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        try:
            tbl.style = "Table Grid"
        except KeyError:
            pass
        _force_full_width(tbl)
        cell = tbl.rows[0].cells[0]
        _shade_cell(cell, "F2F2F2")
        _set_cell_borders(cell, "BFBFBF", sz="4")
        cell.text = ""
        p = cell.paragraphs[0]
        try:
            p.style = doc.styles[PNNL_STYLES["body"]]
        except KeyError:
            pass
        add_inline(p, t)
        doc.add_paragraph()
        return

    header_fill, body_fill, border_color = _CALLOUT_PALETTE[label]

    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    _force_full_width(tbl)

    # --- Header row: colored band, bold white label
    header_cell = tbl.rows[0].cells[0]
    _shade_cell(header_cell, header_fill)
    _set_cell_borders(header_cell, border_color, sz="6")
    header_cell.text = ""
    hp = header_cell.paragraphs[0]
    try:
        hp.style = doc.styles[PNNL_STYLES["body"]]
    except KeyError:
        pass
    # Tighten the header paragraph spacing
    hpPr = hp._p.get_or_add_pPr()
    sp = hpPr.find(qn("w:spacing"))
    if sp is not None:
        hpPr.remove(sp)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "60")
    sp.set(qn("w:after"), "60")
    hpPr.append(sp)
    run = hp.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)

    # --- Body row: tinted, matching border
    body_cell = tbl.rows[1].cells[0]
    _shade_cell(body_cell, body_fill)
    _set_cell_borders(body_cell, border_color, sz="6")
    body_cell.text = ""
    bp = body_cell.paragraphs[0]
    try:
        bp.style = doc.styles[PNNL_STYLES["body"]]
    except KeyError:
        pass
    clean = re.sub(r"^\*\*(WARNING|NOTE|TIP)\.\*\*\s*", "", t)
    add_inline(bp, clean)

    doc.add_paragraph()


# -------- diagram rendering (Word DrawingML primitives) ---------------

# Word DrawingML uses English Metric Units (EMU). 1 inch = 914400 EMU,
# 1 pt = 12700 EMU. Pixels-at-96-dpi → EMU: px * 9525.
EMU_PER_PX = 9525


def _emu(px):
    return int(px * EMU_PER_PX)


# PNNL accent colors. D57500 is the orange used in the cover heading and
# in the PNNL_Solid_Header_Row table style.
PNNL_ORANGE = "D57500"
PNNL_DARK = "1F1F1F"
PNNL_GRAY = "D9D9D9"
PNNL_LIGHT = "F2F2F2"
PNNL_WHITE = "FFFFFF"


def _shape_xml(shape_id, name, x, y, w, h, text_lines, fill, line_color, text_color, bold_first=True):
    """Return the OOXML for a single wps:wsp (rounded rectangle box with text)."""
    # Build paragraphs for each line of text
    text_xml = ""
    for idx, line in enumerate(text_lines):
        bold_attr = "<a:rPr lang=\"en-US\" sz=\"900\" b=\"1\"><a:solidFill><a:srgbClr val=\"{c}\"/></a:solidFill><a:latin typeface=\"Calibri\"/></a:rPr>".format(c=text_color)
        norm_attr = "<a:rPr lang=\"en-US\" sz=\"800\"><a:solidFill><a:srgbClr val=\"{c}\"/></a:solidFill><a:latin typeface=\"Calibri\"/></a:rPr>".format(c=text_color)
        run_attr = bold_attr if (idx == 0 and bold_first) else norm_attr
        text_xml += (
            f'<a:p><a:pPr algn="ctr"/>'
            f'<a:r>{run_attr}<a:t>{_escape(line)}</a:t></a:r>'
            f'</a:p>'
        )

    return f'''
<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <wps:cNvPr id="{shape_id}" name="{name}"/>
  <wps:cNvSpPr/>
  <wps:spPr>
    <a:xfrm>
      <a:off x="{_emu(x)}" y="{_emu(y)}"/>
      <a:ext cx="{_emu(w)}" cy="{_emu(h)}"/>
    </a:xfrm>
    <a:prstGeom prst="roundRect">
      <a:avLst><a:gd name="adj" fmla="val 10000"/></a:avLst>
    </a:prstGeom>
    <a:solidFill><a:srgbClr val="{fill}"/></a:solidFill>
    <a:ln w="9525"><a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill></a:ln>
  </wps:spPr>
  <wps:txbx>
    <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:p>
        <w:pPr><w:jc w:val="center"/><w:spacing w:before="0" w:after="0"/></w:pPr>
        <w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:b/><w:color w:val="{text_color}"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{_escape(text_lines[0])}</w:t></w:r>
      </w:p>
      {"".join(f'<w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="0" w:after="0"/></w:pPr><w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="{text_color}"/><w:sz w:val="16"/></w:rPr><w:t xml:space="preserve">{_escape(line)}</w:t></w:r></w:p>' for line in text_lines[1:])}
    </w:txbxContent>
  </wps:txbx>
  <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="ellipsis" wrap="square" lIns="36000" tIns="18000" rIns="36000" bIns="18000" anchor="ctr" anchorCtr="0"/>
</wps:wsp>
'''


def _connector_xml(shape_id, name, x1, y1, x2, y2, color=PNNL_DARK, arrow=True):
    """Return OOXML for a straight connector with an end arrowhead."""
    # Calculate the bounding box and flipH/flipV
    x = min(x1, x2)
    y = min(y1, y2)
    w = abs(x2 - x1) or 1
    h = abs(y2 - y1) or 1
    flipH = x1 > x2
    flipV = y1 > y2
    flip_attrs = ""
    if flipH:
        flip_attrs += ' flipH="1"'
    if flipV:
        flip_attrs += ' flipV="1"'
    arrow_end = '<a:tailEnd type="triangle" w="med" len="med"/>' if arrow else ""
    return f'''
<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <wps:cNvPr id="{shape_id}" name="{name}"/>
  <wps:cNvSpPr/>
  <wps:spPr>
    <a:xfrm{flip_attrs}>
      <a:off x="{_emu(x)}" y="{_emu(y)}"/>
      <a:ext cx="{_emu(w)}" cy="{_emu(h)}"/>
    </a:xfrm>
    <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
    <a:ln w="12700">
      <a:solidFill><a:srgbClr val="{color}"/></a:solidFill>
      {arrow_end}
    </a:ln>
  </wps:spPr>
  <wps:bodyPr/>
</wps:wsp>
'''


def _label_xml(shape_id, name, x, y, w, h, text, color=PNNL_DARK, italic=False):
    """Free-floating text label (no background)."""
    italic_attr = '<w:i/>' if italic else ''
    return f'''
<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <wps:cNvPr id="{shape_id}" name="{name}"/>
  <wps:cNvSpPr txBox="1"/>
  <wps:spPr>
    <a:xfrm>
      <a:off x="{_emu(x)}" y="{_emu(y)}"/>
      <a:ext cx="{_emu(w)}" cy="{_emu(h)}"/>
    </a:xfrm>
    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
    <a:noFill/>
    <a:ln><a:noFill/></a:ln>
  </wps:spPr>
  <wps:txbx>
    <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:p>
        <w:pPr><w:jc w:val="center"/><w:spacing w:before="0" w:after="0"/></w:pPr>
        <w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>{italic_attr}<w:color w:val="{color}"/><w:sz w:val="16"/></w:rPr><w:t xml:space="preserve">{_escape(text)}</w:t></w:r>
      </w:p>
    </w:txbxContent>
  </wps:txbx>
  <wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr"/>
</wps:wsp>
'''


def _escape(s):
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _topology_diagram_xml():
    """Build the AEMS topology diagram as a wpg group inside a w:drawing.

    Layout (px @ 96 dpi, 600 wide × 480 tall):

        [Operator browser]──►[Traefik proxy 80/443]
                                     │
                ┌────────────────────┼─────────────────────┐
                ▼                    ▼                     ▼
        [Next.js client]       [NestJS API]          [Keycloak SSO]
                                     │                     │
                                     ▼                     ▼
        ┌─────────┬───────────┬─────────────┐      [keycloak-db]
        ▼         ▼           ▼
    [Postgres] [Redis]   [Historian]
                              ▲
                              │
                         [VOLTTRON] ◄── BACnet / RTUs
    """
    W = 620
    H = 520

    # Box coordinates (x, y, w, h, [lines])
    boxes = [
        # id, name, x, y, w, h, lines, fill, line, text
        ("browser", 10, 20, 110, 50, ["Operator", "browser"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("proxy", 230, 10, 180, 70, ["Traefik v3 (proxy)", "TLS termination,", "Host/Path routing"], PNNL_ORANGE, PNNL_ORANGE, PNNL_WHITE),
        ("client", 20, 150, 130, 56, ["Next.js client", "Web UI"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("server", 240, 150, 160, 56, ["NestJS API", "(GraphQL / REST)"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("keycloak", 470, 150, 140, 56, ["Keycloak SSO", "/auth/sso/"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("postgres", 20, 280, 130, 56, ["Postgres", "+ PostGIS"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("redis", 165, 280, 100, 56, ["Redis"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("historian", 280, 280, 130, 56, ["Historian", "(Postgres)"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("kcdb", 470, 280, 140, 56, ["keycloak-db", "(Postgres)"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("volttron", 250, 410, 190, 70, ["VOLTTRON", "+ edge agents", "(BACnet / 47808)"], "404040", "404040", PNNL_WHITE),
    ]

    # Connectors as (x1, y1, x2, y2)
    # Compute by-id box centers for clarity
    cx = {b[0]: b[1] + b[3] / 2 for b in boxes}  # b[1]=x, b[3]=w
    cy = {b[0]: b[2] + b[4] / 2 for b in boxes}  # b[2]=y, b[4]=h
    by_top = {b[0]: b[2] for b in boxes}
    by_bot = {b[0]: b[2] + b[4] for b in boxes}
    bx_left = {b[0]: b[1] for b in boxes}
    bx_right = {b[0]: b[1] + b[3] for b in boxes}

    connectors = [
        # browser → proxy (horizontal arrow)
        (bx_right["browser"], cy["browser"], bx_left["proxy"], cy["browser"]),
        # proxy → client (down-left)
        (cx["proxy"], by_bot["proxy"], cx["client"], by_top["client"]),
        # proxy → server (down)
        (cx["proxy"], by_bot["proxy"], cx["server"], by_top["server"]),
        # proxy → keycloak (down-right)
        (cx["proxy"], by_bot["proxy"], cx["keycloak"], by_top["keycloak"]),
        # server → postgres
        (cx["server"], by_bot["server"], cx["postgres"], by_top["postgres"]),
        # server → redis
        (cx["server"], by_bot["server"], cx["redis"], by_top["redis"]),
        # server → historian
        (cx["server"], by_bot["server"], cx["historian"], by_top["historian"]),
        # keycloak → kcdb
        (cx["keycloak"], by_bot["keycloak"], cx["kcdb"], by_top["kcdb"]),
        # volttron → historian (upward, telemetry)
        (cx["volttron"], by_top["volttron"], cx["historian"], by_bot["historian"]),
    ]

    # Build shape XML
    shapes_xml = []
    sid = 1000
    for box in boxes:
        bid, x, y, w, h, lines, fill, line_c, text_c = box
        shapes_xml.append(
            _shape_xml(sid, bid, x, y, w, h, lines, fill, line_c, text_c)
        )
        sid += 1
    for c in connectors:
        shapes_xml.append(_connector_xml(sid, "conn", *c))
        sid += 1

    # Free-floating labels
    shapes_xml.append(_label_xml(sid, "p_lbl", 420, 5, 90, 18, "80/443"))
    sid += 1
    shapes_xml.append(_label_xml(sid, "rtu_lbl", 450, 425, 160, 40, "◄  RTUs / heat pumps  (BACnet)", italic=True))
    sid += 1

    children = "".join(shapes_xml)

    drawing = f'''
<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <wp:inline distT="0" distB="0" distL="0" distR="0"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
             xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <wp:extent cx="{_emu(W)}" cy="{_emu(H)}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:docPr id="900" name="AEMS Topology"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
        <wpg:wgp>
          <wpg:cNvGrpSpPr/>
          <wpg:grpSpPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{_emu(W)}" cy="{_emu(H)}"/>
              <a:chOff x="0" y="0"/>
              <a:chExt cx="{_emu(W)}" cy="{_emu(H)}"/>
            </a:xfrm>
          </wpg:grpSpPr>
          {children}
        </wpg:wgp>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>
'''
    return drawing


def _wrap_drawing(W, H, doc_name, shapes_xml):
    """Wrap a list of pre-rendered shape XML strings in a single
    wp:inline / wpg:wgp drawing of size (W, H) in pixels.
    """
    children = "".join(shapes_xml)
    return f'''
<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <wp:inline distT="0" distB="0" distL="0" distR="0"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
             xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <wp:extent cx="{_emu(W)}" cy="{_emu(H)}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:docPr id="900" name="{doc_name}"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
        <wpg:wgp>
          <wpg:cNvGrpSpPr/>
          <wpg:grpSpPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{_emu(W)}" cy="{_emu(H)}"/>
              <a:chOff x="0" y="0"/>
              <a:chExt cx="{_emu(W)}" cy="{_emu(H)}"/>
            </a:xfrm>
          </wpg:grpSpPr>
          {children}
        </wpg:wgp>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>
'''


def _os_decision_diagram_xml():
    """OS decision tree:

        ┌──────────────────────┐
        │  What is the host?   │   (orange question box)
        └──────────┬───────────┘
                   │
          ┌────────┼────────┐
          ▼        ▼        ▼
        [PC]    [Pi 5]   [Win 11]
          │        │        │
          ▼        ▼        ▼
    Ubuntu     Ubuntu      Docker
    24.04      24.04 arm64 Desktop
    §4.2       §4.3        §4.4
    """
    W = 620
    H = 360

    question = ("question", 175, 10, 270, 56, ["What is the host?"], PNNL_ORANGE, PNNL_ORANGE, PNNL_WHITE)
    boxes = [
        ("pc", 25, 110, 170, 70, ["Standard x86 PC", "(production)"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("pi", 225, 110, 170, 70, ["Raspberry Pi 5", "(1 building)"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("win", 425, 110, 170, 70, ["Windows workstation", "(dev / demo)"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        # OS / step boxes
        ("pc_os", 25, 240, 170, 88, ["Ubuntu Server", "24.04 LTS (x86_64)", "§4.2"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("pi_os", 225, 240, 170, 88, ["Ubuntu Server", "24.04 LTS arm64", "§4.3 → §4.2"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("win_os", 425, 240, 170, 88, ["Windows 11 +", "Docker Desktop", "§4.4"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
    ]

    shapes = []
    sid = 1100
    # Question box (orange)
    qid, qx, qy, qw, qh, qlines, qfill, qline, qtext = question
    shapes.append(_shape_xml(sid, qid, qx, qy, qw, qh, qlines, qfill, qline, qtext))
    sid += 1
    for b in boxes:
        bid, x, y, w, h, lines, fill, line_c, text_c = b
        shapes.append(_shape_xml(sid, bid, x, y, w, h, lines, fill, line_c, text_c))
        sid += 1

    # Connectors from the question box to each of the three host boxes,
    # and then from each host box down to its OS box.
    q_cx = qx + qw / 2
    q_by = qy + qh
    host_centers = []
    for bid, x, y, w, h, *_ in boxes[:3]:
        host_centers.append((x + w / 2, y, x, x + w, y + h))
    # Question → host (vertical "elbow" via straight lines)
    for cx, top, _, _, _ in host_centers:
        shapes.append(_connector_xml(sid, "conn", q_cx, q_by, cx, top))
        sid += 1
    # Host box bottom → OS box top
    for hidx, host in enumerate(boxes[:3]):
        _, hx, hy, hw, hh, *_ = host
        os = boxes[3 + hidx]
        _, ox, oy, ow, oh, *_ = os
        shapes.append(
            _connector_xml(sid, "conn", hx + hw / 2, hy + hh, ox + ow / 2, oy)
        )
        sid += 1

    return _wrap_drawing(W, H, "OS Decision Tree", shapes)


def _tls_decision_diagram_xml():
    """TLS strategy decision tree:

         ┌────────────────────────────┐
         │ Where will operators reach │   (orange question)
         │   the AEMS UI from?        │
         └─────────────┬──────────────┘
                       │
              ┌────────┼─────────┐
              ▼        ▼         ▼
           Public  Corporate/VPN  Bench/lab
           DNS     internal CA    air-gapped
              │        │              │
              ▼        ▼              ▼
        Let's      Third-party    Self-signed
        Encrypt    cert (drop     (default)
        §6.3       into proxy/)    §6.5
                   §6.4
    """
    W = 640
    H = 360

    question = (
        "question", 165, 10, 310, 60,
        ["Where will operators reach", "the AEMS UI from?"],
        PNNL_ORANGE, PNNL_ORANGE, PNNL_WHITE,
    )
    boxes = [
        ("pub", 25, 110, 180, 70,
         ["Public internet", "with real DNS"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("corp", 230, 110, 180, 70,
         ["Corporate / VPN", "internal CA"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("lab", 435, 110, 180, 70,
         ["Bench / lab /", "air-gapped"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),

        ("le", 25, 240, 180, 88,
         ["Let's Encrypt", "(recommended)", "§6.3"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("third", 230, 240, 180, 88,
         ["Third-party cert", "(drop into proxy/)", "§6.4"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
        ("self", 435, 240, 180, 88,
         ["Self-signed", "(default)", "§6.5"], PNNL_WHITE, PNNL_DARK, PNNL_DARK),
    ]

    shapes = []
    sid = 1200
    qid, qx, qy, qw, qh, qlines, qfill, qline, qtext = question
    shapes.append(_shape_xml(sid, qid, qx, qy, qw, qh, qlines, qfill, qline, qtext))
    sid += 1
    for b in boxes:
        bid, x, y, w, h, lines, fill, line_c, text_c = b
        shapes.append(_shape_xml(sid, bid, x, y, w, h, lines, fill, line_c, text_c))
        sid += 1

    q_cx = qx + qw / 2
    q_by = qy + qh
    # Question → each host-context box
    for bid, x, y, w, h, *_ in boxes[:3]:
        shapes.append(_connector_xml(sid, "conn", q_cx, q_by, x + w / 2, y))
        sid += 1
    # Host-context → matching strategy box
    for hidx in range(3):
        _, hx, hy, hw, hh, *_ = boxes[hidx]
        _, sx, sy, sw, sh, *_ = boxes[3 + hidx]
        shapes.append(_connector_xml(sid, "conn", hx + hw / 2, hy + hh, sx + sw / 2, sy))
        sid += 1

    return _wrap_drawing(W, H, "TLS Decision Tree", shapes)


def _lifecycle_diagram_xml():
    """Deployment lifecycle: four phase boxes connected by arrows.

        ┌──────┐   ┌───────────────┐   ┌───────────────┐   ┌──────────┐
        │ Plan │──►│   Install     │──►│  Site Config  │──►│ Operation│
        │      │   │ (this guide)  │   │ Installer     │   │  Owner   │
        │      │   │               │   │   Guide       │   │  Guide   │
        └──────┘   └───────────────┘   └───────────────┘   └──────────┘
    """
    W = 660
    H = 130

    boxes = [
        # id, x, y, w, h, lines, fill, line, text
        ("plan",    10, 25, 110, 80,
         ["Plan"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("install", 145, 25, 160, 80,
         ["Install", "(this guide)"], PNNL_ORANGE, PNNL_ORANGE, PNNL_WHITE),
        ("site",    330, 25, 160, 80,
         ["Site Config", "Installer Guide"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
        ("op",      515, 25, 130, 80,
         ["Operation", "Owner Guide"], PNNL_LIGHT, PNNL_DARK, PNNL_DARK),
    ]

    shapes = []
    sid = 1300
    for b in boxes:
        bid, x, y, w, h, lines, fill, line_c, text_c = b
        shapes.append(_shape_xml(sid, bid, x, y, w, h, lines, fill, line_c, text_c))
        sid += 1

    # Arrows between successive boxes — horizontal, centered vertically
    for i in range(len(boxes) - 1):
        _, x1, y1, w1, h1, *_ = boxes[i]
        _, x2, y2, _, _, *_ = boxes[i + 1]
        cy = y1 + h1 / 2
        shapes.append(_connector_xml(sid, "conn", x1 + w1, cy, x2, cy))
        sid += 1

    return _wrap_drawing(W, H, "Deployment Lifecycle", shapes)


def render_diagram(doc, name):
    """Render a named diagram as a Word inline drawing."""
    if name == "topology":
        drawing_xml = _topology_diagram_xml()
    elif name == "os_decision":
        drawing_xml = _os_decision_diagram_xml()
    elif name == "tls_decision":
        drawing_xml = _tls_decision_diagram_xml()
    elif name == "lifecycle":
        drawing_xml = _lifecycle_diagram_xml()
    else:
        # Unknown diagram name — fall back to a body-text placeholder so the
        # build does not silently drop content.
        p = doc.add_paragraph(style="Body Text")
        p.add_run(f"[diagram: {name} — not generated]").italic = True
        return

    # Insert a centered paragraph holding the drawing
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Parse the drawing XML and append to the run's element
    from lxml import etree as _etree

    drawing_el = _etree.fromstring(drawing_xml)
    run._r.append(drawing_el)


# -------- post-save zip patching ----------------------------------------


def _patch_customxml_and_headers(docx_path, meta):
    """Patch parts of the .docx zip that python-docx does not write through:

    - customXml/item6.xml stores the PNNL cover bindings (title, authors,
      date, number). The visible cover-page paragraphs were patched in
      _patch_title_page, but Word can re-render the bound content controls
      from this customXml when fields are refreshed, so the values here
      must match too.
    - word/header5.xml is a left-over template header that prints
      'PNNL-35747' at the top of certain page-section ranges. Replace with
      the new report number so the running head is consistent.
    """
    import zipfile
    import tempfile
    import os

    title = meta.get("title", "AEMS Software Deployment Guide")
    subtitle = meta.get("subtitle", "")
    report_number = meta.get("report_number", "PNNL-XXXXX")
    pub_date_iso = meta.get("date_iso", "2026-06-01T00:00:00")
    authors_raw = meta.get("authors", "Pacific Northwest National Laboratory")
    authors = authors_raw.replace("\\n", "\n")

    # New customXml/item6.xml contents — preserves the PNNL_Template
    # namespace + element structure expected by the cover content controls.
    new_item6 = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<projectDoc xmlns="PNNL_Template">\n'
        '  <PNNL_Template>\n'
        '    <PNNL_Content_Controls>\n'
        f'      <PNNL_Contract>{_xml_escape(meta.get("contract", ""))}</PNNL_Contract>\n'
        '      <PNNL_Designation/>\n'
        f'      <PNNL_Title>{_xml_escape(title)}</PNNL_Title>\n'
        f'      <PNNL_Subtitle>{_xml_escape(subtitle)}</PNNL_Subtitle>\n'
        f'      <PNNL_Date>{_xml_escape(pub_date_iso)}</PNNL_Date>\n'
        f'      <PNNL_Authors>{_xml_escape(authors)}</PNNL_Authors>\n'
        f'      <PNNL_Number>{_xml_escape(report_number)}</PNNL_Number>\n'
        '    </PNNL_Content_Controls>\n'
        '  </PNNL_Template>\n'
        '</projectDoc>'
    )

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "customXml/item6.xml":
                    data = new_item6.encode("utf-8")
                elif item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                    text = data.decode("utf-8", errors="replace")
                    if "PNNL-35747" in text:
                        text = text.replace("PNNL-35747", report_number)
                        data = text.encode("utf-8")
                zout.writestr(item, data)
        os.replace(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def _xml_escape(s):
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


# -------- main ----------------------------------------------------------


def _prewalk_section_numbers(blocks):
    """Populate SECTION_BOOKMARKS by walking the parsed blocks without
    rendering. Same numbering rules as render_heading so forward
    references like '§4' resolve correctly when seen in the Preface
    before §4 is actually rendered.
    """
    SECTION_BOOKMARKS.clear()
    h1 = h2 = h3 = 0
    in_unnumbered_h1 = False
    for kind, payload in blocks:
        if kind not in ("h1", "h2", "h3"):
            continue
        text = payload.strip()
        norm = text.lower()
        if kind == "h1":
            if norm in FRONT_HEADINGS or norm.startswith("appendix "):
                in_unnumbered_h1 = True
                continue
            in_unnumbered_h1 = False
            h1 += 1
            h2 = 0
            h3 = 0
            num = str(h1)
        elif kind == "h2":
            if in_unnumbered_h1 or norm in FRONT_HEADINGS:
                continue
            h2 += 1
            h3 = 0
            num = f"{h1}.{h2}"
        else:  # h3
            if in_unnumbered_h1:
                continue
            h3 += 1
            num = f"{h1}.{h2}.{h3}"
        SECTION_BOOKMARKS[num] = _bookmark_for_section(num)


def _insert_toc(doc):
    """Insert a Word TOC field. Word renders the populated TOC on first
    open via 'Update Field' or after the reader presses F9. The
    placeholder text is shown until that happens.
    """
    # Heading line for the TOC
    p_heading = doc.add_paragraph(style=PNNL_STYLES["front_h1"])
    p_heading.add_run("Contents")

    # The TOC field: a w:fldSimple with the appropriate instruction.
    # \o "1-3" includes headings 1-3; \h makes entries hyperlinks; \z hides
    # tab leader chars in web layout; \u uses applied paragraph outline
    # level (so our Heading 1/2/3 styles are picked up).
    p_toc = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r' TOC \o "1-3" \h \z \u ')
    placeholder_run = OxmlElement("w:r")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.set(qn("xml:space"), "preserve")
    placeholder_t.text = (
        "Right-click and choose Update Field to populate the table of contents."
    )
    placeholder_run.append(placeholder_t)
    fld.append(placeholder_run)
    p_toc._p.append(fld)

    # Page break so body content starts on a fresh page
    p_br = doc.add_paragraph()
    run = p_br.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def build():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()
    meta, body_lines = parse_frontmatter(lines)

    blocks = list(parse_blocks(body_lines))

    # Pre-walk: assign section numbers so forward references resolve.
    _prewalk_section_numbers(blocks)

    # Reset render-time state (incremented again in render_heading)
    _SECTION_STATE["h1"] = 0
    _SECTION_STATE["h2"] = 0
    _SECTION_STATE["h3"] = 0
    _SECTION_STATE["in_unnumbered_h1"] = False

    doc = clone_template_and_clear_body(meta)

    # Insert the TOC field before any body content. Word will render the
    # populated table on first open / Update Field.
    _insert_toc(doc)

    for kind, payload in blocks:
        if kind in ("h1", "h2", "h3", "h4"):
            render_heading(doc, int(kind[1]), payload)
        elif kind == "p":
            render_body(doc, payload)
        elif kind == "code":
            render_code(doc, payload)
        elif kind == "table":
            render_table(doc, *payload)
        elif kind == "ul":
            render_list(doc, payload, "ul")
        elif kind == "ol":
            render_list(doc, payload, "ol")
        elif kind == "hr":
            doc.add_paragraph()
        elif kind == "blockquote":
            render_blockquote(doc, payload)
        elif kind == "diagram":
            render_diagram(doc, payload)

    doc.save(str(OUT))
    _patch_customxml_and_headers(str(OUT), meta)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
